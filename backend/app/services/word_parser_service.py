from typing import Dict, List
import docx
from docx.document import Document
from docx.table import Table, _Cell
import xml.etree.ElementTree as ET

class WordParserService:
    def __init__(self):
        # python-docx is used for parsing Word documents
        pass

    def extract_text_content(self, file_path: str) -> Dict:
        """
        Extracts text content, structure (paragraph styles), and metadata from a Word document.
        """
        try:
            document: Document = docx.Document(file_path)

            full_text = [para.text for para in document.paragraphs]

            structure_info = [
                {
                    "text_preview": para.text[:100] + "..." if len(para.text) > 100 else para.text,
                    "style": para.style.name if para.style else "Normal",
                    "is_heading": para.style.name.startswith("Heading") if para.style else False,
                    "outline_level": para.paragraph_format.outline_level
                                     if para.paragraph_format.outline_level is not None else -1 # -1 if not set
                }
                for para in document.paragraphs
            ]

            core_props = document.core_properties
            metadata = {
                "author": core_props.author,
                "category": core_props.category,
                "comments": core_props.comments,
                "content_status": core_props.content_status,
                "created": core_props.created.isoformat() if core_props.created else None,
                "identifier": core_props.identifier,
                "keywords": core_props.keywords,
                "language": core_props.language,
                "last_modified_by": core_props.last_modified_by,
                "last_printed": core_props.last_printed.isoformat() if core_props.last_printed else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "revision": core_props.revision,
                "subject": core_props.subject,
                "title": core_props.title,
                "version": core_props.version,
            }

            return {
                "text": "\n".join(full_text),
                "structure": structure_info,
                "metadata": metadata
            }
        except Exception as e:
            return {
                "error": f"Failed to extract text content: {str(e)}",
                "text": "",
                "structure": [],
                "metadata": {}
            }

    def extract_tables(self, file_path: str) -> List[Dict]:
        """
        Extracts tables from a Word document.
        Each table is represented as a list of lists (rows and cells).
        """
        tables_data = []
        try:
            document: Document = docx.Document(file_path)
            for i, table_obj in enumerate(document.tables):
                table_content = []
                current_table: Table = table_obj
                for row_idx, row in enumerate(current_table.rows):
                    row_content = []
                    cell: _Cell
                    for col_idx, cell in enumerate(row.cells):
                        # Basic handling for merged cells (first cell in a merge retains content)
                        # More complex merge analysis would require checking cell.merged_cells
                        row_content.append(cell.text)
                    table_content.append(row_content)

                # Try to find a caption for the table (common patterns)
                caption_text = "No caption found"
                # Heuristic: Look for a paragraph immediately before or after the table
                # This requires knowing the table's position, which is hard with python-docx alone.
                # For now, this part is omitted as it's unreliable without element context.

                tables_data.append({
                    "table_index": i,
                    "rows": len(table_content),
                    "columns": len(table_content[0]) if table_content else 0,
                    "data": table_content,
                    "caption_guess": caption_text
                })
        except Exception as e:
            return [{"error": f"Failed to extract tables: {str(e)}"}]
        return tables_data

    def extract_headers_and_sections(self, file_path: str) -> Dict:
        """
        Extracts document chapter structure based on heading styles or outline levels.
        Returns a hierarchical structure of sections.
        """
        # Using outline_level is generally more reliable than style names.
        # Word outline levels: 0-8 correspond to Heading 1-9.
        # Level 9 means body text (not an outline level).

        # This function aims to build a nested dictionary representing the hierarchy.
        # Example: {"Heading 1 Text": {"level": 1, "paragraphs": [], "subheadings": {"Subheading 1.1": ...}}}

        root_sections = {} # Top-level sections
        path_to_current_section = [] # Tracks parent headings texts for nesting

        try:
            document: Document = docx.Document(file_path)

            # Add a default "Prologue" or "Introduction" for content before the first heading
            # This will hold content until the first actual heading is encountered.
            # We'll add it to root_sections and make it the initial current_section.

            current_section_dict = {"level": 0, "paragraphs": [], "subheadings": {}}
            # We need a way to point to where the current section's content should go.
            # Let's maintain a reference to the dictionary of the current heading.

            # To handle content before any heading:
            content_before_first_heading = []
            first_heading_found = False

            for para in document.paragraphs:
                para_text = para.text.strip()
                outline_level = para.paragraph_format.outline_level # 0-8 for Headings 1-9, None for others

                is_heading = outline_level is not None and outline_level < 9 # Max Word outline level for headings

                if is_heading:
                    first_heading_found = True
                    level = outline_level + 1 # Convert 0-8 to 1-9 for user display

                    # Adjust path: pop headings from path if current level is same or higher
                    while len(path_to_current_section) >= level:
                        path_to_current_section.pop()

                    # Find parent dictionary for this heading
                    parent_dict = root_sections
                    for heading_text_in_path in path_to_current_section:
                        # It's possible a heading was added without subheadings dict if it had no children yet
                        if "subheadings" not in parent_dict[heading_text_in_path]:
                             parent_dict[heading_text_in_path]["subheadings"] = {}
                        parent_dict = parent_dict[heading_text_in_path]["subheadings"]

                    # Add current heading to its parent
                    # Handle duplicate heading names at the same level if necessary (e.g., by appending count)
                    # For now, assuming unique heading names within the same parent/level for simplicity
                    if not para_text: # Use a placeholder if heading text is empty
                        para_text = f"Untitled Heading (Level {level})"

                    # Ensure no overwrite of existing heading with same name at same level by making key unique if needed
                    original_para_text = para_text
                    counter = 1
                    while para_text in parent_dict:
                        para_text = f"{original_para_text} ({counter})"
                        counter += 1

                    parent_dict[para_text] = {"level": level, "paragraphs": [], "subheadings": {}}
                    path_to_current_section.append(para_text)

                else: # Not a heading, it's a paragraph of content
                    if not first_heading_found:
                        content_before_first_heading.append(para_text)
                    else:
                        # Add paragraph to the current section (deepest in path)
                        current_section_parent_dict = root_sections
                        for heading_text_in_path in path_to_current_section[:-1]: # Navigate to parent of current
                            current_section_parent_dict = current_section_parent_dict[heading_text_in_path]["subheadings"]

                        if path_to_current_section: # If there is any heading active
                            current_heading_text = path_to_current_section[-1]
                            current_section_parent_dict[current_heading_text]["paragraphs"].append(para_text)
                        else: # Should not happen if first_heading_found is true, but as fallback
                            content_before_first_heading.append(para_text)

            if content_before_first_heading:
                 # Prepend an "Introduction" or similar section for this content
                 # Needs to be handled carefully if root_sections already has items due to later processing
                 # A simple way:
                 intro_key = "Introduction / Uncategorized Content"
                 if intro_key in root_sections: # If it was somehow created
                     root_sections[intro_key]["paragraphs"] = content_before_first_heading + root_sections[intro_key].get("paragraphs", [])
                 else: # Prepend it
                     updated_root_sections = {intro_key: {"level": 0, "paragraphs": content_before_first_heading, "subheadings": {}}}
                     updated_root_sections.update(root_sections) # Add existing sections after
                     root_sections = updated_root_sections


            return root_sections

        except Exception as e:
            return {"error": f"Failed to extract headers and sections: {str(e)}"}


    def extract_images_info(self, file_path: str) -> List[Dict]:
        """
        Extracts information about images, attempting to link them to rIds and find captions.
        """
        images_info = []
        try:
            document: Document = docx.Document(file_path)

            # Map rIds to image parts (filename, content_type)
            img_rels = {}
            for r_id, rel in document.part.rels.items():
                if "image" in rel.target_ref: # rel.target_ref is like '../media/image1.png'
                    img_part = rel.target_part
                    img_rels[r_id] = {
                        "partname": img_part.partname, # e.g., '/word/media/image1.png'
                        "content_type": img_part.content_type # e.g., 'image/png'
                    }

            image_idx_counter = 0 # Global counter for images found

            # Iterate through paragraphs to find <w:drawing> elements which host images
            for para_idx, para in enumerate(document.paragraphs):
                # The XML structure for an image is typically:
                # w:p -> w:r -> w:drawing -> wp:inline (or wp:anchor) -> a:graphic -> a:graphicData -> pic:pic -> ... -> a:blip (with r:embed)
                # Namespaces:
                # xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                # xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                # xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                # xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                # xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"

                # Define namespaces for ET.fromstring().find()
                ns = {
                    'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                    'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
                    'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture",
                    'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                }

                try:
                    # Parse the paragraph's XML structure
                    p_element = ET.fromstring(para._p.xml) # _p is the lxml element for paragraph

                    # Find all <a:blip> elements with an r:embed attribute within this paragraph
                    # These are the actual image references.
                    for blip_element in p_element.findall('.//a:blip', namespaces=ns):
                        rId = blip_element.get(f'{{{ns["r"]}}}embed')
                        if rId and rId in img_rels:
                            img_data = img_rels[rId]

                            # Attempt to find caption
                            caption_text = "No caption found"
                            # Heuristic 1: Paragraph immediately following, styled as "Caption" or starts with "Figure", "Fig.", "图"
                            if para_idx + 1 < len(document.paragraphs):
                                next_para = document.paragraphs[para_idx+1]
                                if (next_para.style and "caption" in next_para.style.name.lower()) or \
                                   next_para.text.strip().lower().startswith(("figure", "fig.", "图")):
                                    caption_text = next_para.text.strip()

                            # Heuristic 2: Paragraph immediately preceding (less common for figures)
                            if caption_text == "No caption found" and para_idx > 0:
                                prev_para = document.paragraphs[para_idx-1]
                                if (prev_para.style and "caption" in prev_para.style.name.lower()) or \
                                   prev_para.text.strip().lower().startswith(("figure", "fig.", "图")): # Less likely for previous
                                    caption_text = prev_para.text.strip()

                            # Try to get image name/description from <wp:docPr> if available
                            # <wp:inline> or <wp:anchor> -> <wp:docPr id="1" name="Picture 1" descr="Some description">
                            image_name_prop = "N/A"
                            image_descr_prop = "N/A"

                            # Find parent <wp:inline> or <wp:anchor> to search for <wp:docPr>
                            # This requires navigating up from blip_element, which is complex with ET.
                            # Alternatively, find all docPr elements and see if one is an ancestor.
                            # For simplicity, we'll focus on what's directly available or easier.

                            images_info.append({
                                "image_index": image_idx_counter,
                                "internal_rId": rId,
                                "filename_in_docx": img_data["partname"].split('/')[-1],
                                "content_type": img_data["content_type"],
                                "paragraph_index_approx": para_idx, # Paragraph where drawing element is found
                                "caption_text_guess": caption_text,
                                # "name_property": image_name_prop, # Would require more XML parsing
                                # "description_property": image_descr_prop, # Would require more XML parsing
                            })
                            image_idx_counter += 1
                except ET.ParseError:
                    # Issue parsing this paragraph's XML, skip.
                    # print(f"Warning: XML ParseError for paragraph {para_idx}")
                    pass
                except Exception as xml_e:
                    # Other XML related issues
                    # print(f"Warning: XML processing error for paragraph {para_idx}: {xml_e}")
                    pass

        except Exception as e:
            return [{"error": f"Failed to extract images info: {str(e)}"}]
        return images_info

# Main block for local testing (commented out for agent use)
# if __name__ == '__main__':
#     from docx import Document as DocCreator
#     from docx.shared import Inches, Pt
#     from docx.enum.style import WD_STYLE_TYPE

#     print("Creating a dummy test_document.docx for testing WordParserService...")
#     try:
#         doc = DocCreator()

#         # Set up styles for testing outline levels
#         styles = doc.styles
#         try:
#             styles.add_style('TestHeading1', WD_STYLE_TYPE.PARAGRAPH)
#             styles['TestHeading1'].font.name = 'Calibri'
#             styles['TestHeading1'].font.size = Pt(16)
#             styles['TestHeading1'].paragraph_format.outline_level = 0 # Heading 1

#             styles.add_style('TestHeading2', WD_STYLE_TYPE.PARAGRAPH)
#             styles['TestHeading2'].font.name = 'Calibri'
#             styles['TestHeading2'].font.size = Pt(14)
#             styles['TestHeading2'].paragraph_format.outline_level = 1 # Heading 2

#             styles.add_style('TestCaption', WD_STYLE_TYPE.PARAGRAPH)
#             styles['TestCaption'].font.italic = True
#         except Exception as style_e:
#             print(f"Could not create custom styles (may exist if run multiple times): {style_e}")


#         doc.add_paragraph('Content before any heading.')
#         doc.add_paragraph('More content before any heading.')

#         doc.add_paragraph('Main Title (Styled as H1)', style='TestHeading1')
#         doc.add_paragraph('This is a paragraph under Main Title.')

#         doc.add_paragraph('Subtitle 1.1 (Styled as H2)', style='TestHeading2')
#         doc.add_paragraph('Paragraph under Subtitle 1.1.')
#         doc.add_paragraph('Another paragraph under Subtitle 1.1.')

#         # Add a table
#         doc.add_paragraph("Table below should have a caption attempt.")
#         table_data = [["A", "B"], ["1", "2"]]
#         table = doc.add_table(rows=2, cols=2)
#         table.cell(0,0).text = table_data[0][0]
#         table.cell(0,1).text = table_data[0][1]
#         table.cell(1,0).text = table_data[1][0]
#         table.cell(1,1).text = table_data[1][1]
#         doc.add_paragraph("Table 1: A sample table.", style='TestCaption')


#         doc.add_paragraph('Subtitle 1.2 (Styled as H2)', style='TestHeading2')
#         doc.add_paragraph('Paragraph under Subtitle 1.2.')

#         doc.add_paragraph('Another Main Title (Styled as H1)', style='TestHeading1')
#         doc.add_paragraph('Paragraph under the second Main Title.')

#         # Image handling: requires a placeholder image. Create one if it doesn't exist.
#         placeholder_image_path = 'placeholder_image.png'
#         # try:
#         #     from PIL import Image, ImageDraw
#         #     img = Image.new('RGB', (100, 50), color = 'red')
#         #     d = ImageDraw.Draw(img)
#         #     d.text((10,10), "Test IMG", fill=(255,255,0))
#         #     img.save(placeholder_image_path)
#         #     print(f"Created '{placeholder_image_path}' for testing.")
#         #     doc.add_picture(placeholder_image_path, width=Inches(1))
#         #     doc.add_paragraph("Figure 1: A test image.", style='TestCaption')
#         # except ImportError:
#         #     print("Pillow not installed, cannot create/add placeholder image. Skipping image test.")
#         # except FileNotFoundError:
#         #      print(f"'{placeholder_image_path}' not found and Pillow not available to create it.")
#         # except Exception as img_e:
#         #     print(f"Error adding image: {img_e}")


#         doc.save('test_document.docx')
#         print("Dummy 'test_document.docx' created/updated.")

#         parser = WordParserService()
#         file_path = 'test_document.docx'

#         print(f"\n--- Extracting Text Content from {file_path} ---")
#         text_content = parser.extract_text_content(file_path)
#         if "error" in text_content: print(f"Error: {text_content['error']}")
#         else:
#             print("Text (first 200 chars):", text_content['text'][:200] + "...")
#             # print("Structure (first 3 items):", text_content['structure'][:3])
#             print("Metadata (title):", text_content['metadata'].get('title'))

#         print(f"\n--- Extracting Tables from {file_path} ---")
#         tables = parser.extract_tables(file_path)
#         if tables and isinstance(tables, list) and tables[0].get("error"): print(f"Error: {tables[0]['error']}")
#         else:
#             for idx, table_info in enumerate(tables):
#                 print(f"Table {idx}: {table_info['rows']}x{table_info['columns']}, Caption: {table_info.get('caption_guess')}")

#         print(f"\n--- Extracting Headers and Sections from {file_path} ---")
#         sections = parser.extract_headers_and_sections(file_path)
#         if isinstance(sections, dict) and sections.get("error"): print(f"Error: {sections['error']}")
#         else:
#             import json
#             print(json.dumps(sections, indent=2))


#         print(f"\n--- Extracting Images Info from {file_path} ---")
#         images = parser.extract_images_info(file_path)
#         if images and isinstance(images, list) and images[0].get("error"): print(f"Error: {images[0]['error']}")
#         else:
#             if not images: print("No images found or extracted.")
#             for img_info in images: print(img_info)

#         # import os
#         # if os.path.exists(placeholder_image_path): os.remove(placeholder_image_path)
#         # if os.path.exists('test_document.docx'): os.remove('test_document.docx')
#         # print("\nCleaned up dummy files.")

#     except Exception as e:
#         print(f"An error occurred during the test setup or execution: {e}")
#         import traceback
#         traceback.print_exc()
